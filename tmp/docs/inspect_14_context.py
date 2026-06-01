# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
path = Path(r'C:/Users/94122/Desktop/毕设/1/out/heart_sensor_version1.1_backup_20260413_section_2_3_3.docx')
doc = Document(str(path))
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('1.4') or t.startswith('1.3') or t.startswith('1.5'):
        start = max(0, i-6)
        end = min(len(doc.paragraphs), i+12)
        for j in range(start, end):
            print(f'{j}: {doc.paragraphs[j].text}')
        print('---')
