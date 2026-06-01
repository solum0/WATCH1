# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

text_blocks = [
    '人机交互单元主要用于完成模式切换、时间设置、闹钟设置以及功能启停等操作，是用户对健康运动监测仪进行控制的重要输入部分。考虑到本设计采用手表式结构，交互器件不仅要具备体积小、功耗低、操作直观和响应迅速等特点，还应尽量减少外围电路复杂度，便于与 STM32 主控芯片进行连接和程序实现。因此，在器件选择过程中重点考虑按键数量、安装方式、可靠性以及对系统 I/O 资源的占用情况。',
    '综合比较后，本设计选用独立轻触按键作为人机交互单元的主要输入器件。与矩阵键盘相比，独立按键结构简单、成本较低、布线方便，更适合本系统功能数量较少但要求响应及时的可穿戴应用场景。结合工程程序可以看出，系统采用外部中断方式采集按键事件，并在软件中完成模式切换、时间与闹钟参数设置、秒表控制等功能处理，这种实现方式能够减少主控芯片持续轮询带来的资源占用，提高交互响应速度，同时也有利于系统的低功耗运行。',
    '此外，轻触按键具有机械反馈明确、操作手感清晰、使用寿命较长等优点，能够满足日常佩戴条件下频繁操作的需求。其体积小、封装灵活，便于在有限的 PCB 空间内完成布局，并可与 3.3 V 逻辑电平系统直接配合，硬件实现难度较低。综上所述，独立轻触按键在体积、成本、可靠性及软件实现复杂度等方面均符合本设计需求，因此最终选用其作为健康运动监测仪的人机交互输入器件，其实物如图 2.3 所示。'
]

src = sorted(Path(r'C:/Users/94122/Desktop').glob('*/1/out/heart_sensor_version1.1.backup_20260413_*.docx'))[0]
out = src.with_name('heart_sensor_version1.1_backup_20260413_section_2_3_3.docx')
doc = Document(str(src))
heading_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith('2.3.3'))
image_para = doc.paragraphs[heading_idx + 1]
sample_para = doc.paragraphs[131]

for block in reversed(text_blocks):
    new_p = image_para.insert_paragraph_before(block, style=sample_para.style)
    new_p.paragraph_format.first_line_indent = sample_para.paragraph_format.first_line_indent
    new_p.paragraph_format.left_indent = sample_para.paragraph_format.left_indent
    new_p.paragraph_format.right_indent = sample_para.paragraph_format.right_indent
    new_p.paragraph_format.space_before = sample_para.paragraph_format.space_before
    new_p.paragraph_format.space_after = sample_para.paragraph_format.space_after
    new_p.paragraph_format.line_spacing = sample_para.paragraph_format.line_spacing
    if sample_para.runs:
        sample_run = sample_para.runs[0]
        for run in new_p.runs:
            run.bold = False
            run.italic = False
            run.underline = False
            run.font.size = sample_run.font.size
            if sample_run.font.name:
                run.font.name = sample_run.font.name

doc.save(str(out))
print(out)
