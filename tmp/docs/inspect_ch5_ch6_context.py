# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
path = Path(r'C:/Users/94122/Desktop/毕设/1/out/heart_sensor_version1.1_backup_20260413_section_2_3_3.docx')
doc = Document(str(path))
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('5 ') or p.text.strip().startswith('6 '):
        for j in range(i, min(i+10, len(doc.paragraphs))):
            print(f'{j}: {doc.paragraphs[j].text}')
        print('---')
