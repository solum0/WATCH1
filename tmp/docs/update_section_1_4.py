# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

text_blocks = [
    '为了便于对本文研究内容与实现过程进行系统阐述，全文共分为六章，各章内容安排如下：',
    '（1）第 1 章为绪论。主要介绍课题研究背景、国内外研究现状、本文主要研究内容以及论文整体结构安排，为后续章节展开奠定基础。',
    '（2）第 2 章为总体设计方案。结合健康运动监测仪的功能需求，完成系统总体架构设计，并对主控、显示、人机交互、信息采集、实时时钟、提醒以及 BLE 通信等关键器件进行选型分析。',
    '（3）第 3 章为硬件设计。重点介绍系统硬件总体结构以及主控、电源、人机交互、传感器采集、实时时钟和蓝牙通信等电路设计内容，并对 PCB 实现方案进行说明。',
    '（4）第 4 章为系统程序设计。围绕基于 FreeRTOS 的软件架构，介绍开发工具、主程序流程及各功能子程序流程设计，说明系统软件实现思路与任务协同机制。',
    '（5）第 5 章为系统实现与测试。对样机的软硬件实现情况进行说明，并通过模式切换、时间设置、闹钟提醒、温湿度采集、秒表计时、心率检测、计步统计和 BLE 控制等测试验证系统功能与稳定性。',
    '（6）第 6 章为结语。对本文完成的主要研究工作进行总结，分析系统存在的不足，并对后续优化方向进行展望。',
    ''
]

path = Path(r'C:/Users/94122/Desktop/毕设/1/out/heart_sensor_version1.1_backup_20260413_section_2_3_3.docx')
out_path = path
fallback_path = path.with_name('heart_sensor_version1.1_backup_20260413_section_2_3_3_v2.docx')
doc = Document(str(path))
heading_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == '1.4 文章设计结构')
next_heading_idx = next(i for i, p in enumerate(doc.paragraphs) if i > heading_idx and p.text.strip().startswith('2 '))
sample_para = doc.paragraphs[heading_idx + 1]

# Remove existing content between 1.4 and chapter 2 heading.
for idx in range(next_heading_idx - 1, heading_idx, -1):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

next_heading_para = next(p for p in doc.paragraphs if p.text.strip().startswith('2 '))
for block in reversed(text_blocks):
    new_p = next_heading_para.insert_paragraph_before(block, style=sample_para.style)
    new_p.paragraph_format.first_line_indent = sample_para.paragraph_format.first_line_indent
    new_p.paragraph_format.left_indent = sample_para.paragraph_format.left_indent
    new_p.paragraph_format.right_indent = sample_para.paragraph_format.right_indent
    new_p.paragraph_format.space_before = sample_para.paragraph_format.space_before
    new_p.paragraph_format.space_after = sample_para.paragraph_format.space_after
    new_p.paragraph_format.line_spacing = sample_para.paragraph_format.line_spacing
    if sample_para.runs:
        sample_run = sample_para.runs[0]
        for run in new_p.runs:
            run.bold = False
            run.italic = False
            run.underline = False
            run.font.size = sample_run.font.size
            if sample_run.font.name:
                run.font.name = sample_run.font.name

saved_to = None
try:
    doc.save(str(out_path))
    saved_to = out_path
except PermissionError:
    doc.save(str(fallback_path))
    saved_to = fallback_path

print(saved_to)
